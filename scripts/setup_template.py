import argparse
import os
from docx import Document

parser = argparse.ArgumentParser(description="Clean up source document to create a base template.")
parser.add_argument('--source', type=str, default='collections/source_template.docx', help='Path to the source docx file')
parser.add_argument('--target', type=str, default='collections/base_template.docx', help='Path to save the cleaned base docx file')
args = parser.parse_args()

source_path = args.source
target_path = args.target

if not os.path.exists(source_path):
    print(f"Error: Source file '{source_path}' does not exist.")
    exit(1)

doc = Document(source_path)

# Delete all tables
for table in reversed(doc.tables):
    table._element.getparent().remove(table._element)

for p in list(doc.paragraphs):
    text = p.text.strip().lower()
    should_delete = False
    
    if text == "daftar revisi" or text == "dokumen hasil testing" or "cek usia -1 hari" in text or "nomor transaksi" in text or "tanggal_" in text or "tenor" in text:
        should_delete = True
    elif any(keyword in text for keyword in ["tc-", "test case", "precondition", "test data", "status", "tabel"]):
        should_delete = True
        
    if should_delete:
        try:
            p._element.getparent().remove(p._element)
        except Exception:
            pass


doc.save(target_path)
print("Base template fully cleaned at", target_path)
