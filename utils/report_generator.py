import os
import copy
import json
import traceback
from datetime import datetime
from jinja2 import Environment, FileSystemLoader

try:
    from weasyprint import HTML
    HAS_WEASYPRINT = True
except (ImportError, OSError):
    HAS_WEASYPRINT = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.shared import OxmlElement as SharedOxmlElement
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from openpyxl import load_workbook
    from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = SharedOxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t')
    t.text = "1"
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(t)
    run._r.append(fldChar3)

class ReportGenerator:
    def __init__(self, template_dir="templates"):
        self.env = Environment(loader=FileSystemLoader(template_dir))

    def _filter_db_results(self, evidences: dict) -> dict:
        evidences_copy = copy.deepcopy(evidences)
        allowed_cols = {
            'nomor_transaksi', 'status_akseptasi', 'tanggal_lahir', 
            'tanggal_rencana_realisasi', 'tanggal_akhir_asuransi', 
            'tenor', 'uang_pertanggungan', 'nomor_loan',
            'ktp', 'id_debitur', 'nilai_pertanggungan', 'nama_debitur', 'Validasi DB', 'Response API', 'Status Code', 'Status Akseptasi',
            'kode_bank', 'jenis_covering', 'jangka_waktu', 'nomor_rekening_pinjaman', 'nomor_perjanjian_kredit', 'tanggal_mulai_covering', 'premi',
            'id_sertifikat', 'no_sertifikat', 'tgl_sertifikat', 'url_download_sertifikat', 'is_polis_sent'
        }
        for tc_id, data in evidences_copy.items():
            for db in data.get('db', []):
                if db.get('result') and isinstance(db['result'], list) and len(db['result']) > 0:
                    if isinstance(db['result'][0], dict):
                        new_results = []
                        for row in db['result']:
                            filtered_row = {k: v for k, v in row.items() if str(k).lower() in allowed_cols}
                            new_results.append(filtered_row if filtered_row else row)
                        db['result'] = new_results
            
            # Tambahkan actual_result agar bisa dipakai di HTML dan DOCX
            actual_json = {}
            if data.get("api") and len(data["api"]) > 0:
                actual_json = data["api"][-1].get("response_json", {})
            if actual_json:
                expected = data.get("expected_result", "")
                status = data.get("status", "").lower()
                
                db_msg = ""
                if "tidak kerecord" in expected.lower():
                    db_msg = "\n\n2. data tidak kerecord di DB"
                else:
                    if status != "failed":
                        db_msg = "\n\n2. data DB sesuai dengan request postman"
                
                data["actual_result"] = "1. Response API:\n" + json.dumps(actual_json, indent=2) + db_msg
            else:
                data["actual_result"] = ""

        return evidences_copy

    def generate_html(self, evidences: dict, output_file: str):
        filtered_evidences = self._filter_db_results(evidences)
        template = self.env.get_template("report_template.html")
        current_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        html_out = template.render(
            evidences=filtered_evidences,
            current_date=current_date
        )
        
        os.makedirs(os.path.dirname(output_file), exist_ok=True)
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(html_out)
        return html_out

    def generate_pdf(self, evidences: dict, output_file: str):
        html_file = output_file.replace(".pdf", ".html")
        self.generate_html(evidences, html_file)
        
        try:
            HTML(html_file).write_pdf(output_file)
            print(f"PDF generated successfully: {output_file}")
        except ImportError:
            print("WeasyPrint is not installed. Skipping PDF generation, HTML saved instead.")
        except Exception as e:
            print(f"Failed to generate PDF: {e}")

    def generate_docx(self, evidences: dict, output_file: str):
        filtered_evidences = self._filter_db_results(evidences)
        try:
            
            # Use base_template
            doc = Document('collections/base_template.docx')
            
            # Try to get normal style, fallback if issue
            try:
                style = doc.styles['Normal']
                font = style.font
                font.name = 'Arial'
                font.size = Pt(10)
            except:
                pass
            
            current_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            current_date_cover = datetime.now().strftime("%d %B %Y")
            
            report_title = os.environ.get("REPORT_TITLE", "Laporan Hasil Pengujian API")
            project_code = os.environ.get("PROJECT_CODE", "PRJ-000")
            
            for p in doc.paragraphs:
                # Cek jika REPORT_TITLE dan PROJECT_CODE ada di paragraf yang sama
                if '{{REPORT_TITLE}}' in p.text and ('{{PROJECT_CODE}}' in p.text or '{{Project Code}}' in p.text):
                    p.text = "" # Bersihkan teks paragraf
                    
                    # Run 1: Report Title
                    run_title = p.add_run(report_title)
                    run_title.font.name = 'Tahoma'
                    run_title.font.size = Pt(20)
                    run_title.font.bold = True
                    
                    # Line Break dengan jarak terkontrol agar tidak terlalu dempet
                    run_gap = p.add_run('\n\n')
                    run_gap.font.size = Pt(8)
                    
                    # Run 2: Project Code
                    run_code = p.add_run(project_code)
                    run_code.font.name = 'Tahoma'
                    run_code.font.size = Pt(15)
                    run_code.font.bold = True
                    
                    # Kembalikan line spacing ke normal agar jika judul wrap ke bawah tidak renggang
                    p.paragraph_format.line_spacing = 1.5
                    continue
                    
                is_report_title = False
                is_project_code = False
                is_date = False

                if '{{REPORT_TITLE}}' in p.text:
                    p.text = p.text.replace('{{REPORT_TITLE}}', report_title)
                    is_report_title = True
                if '{{PROJECT_CODE}}' in p.text or '{{Project Code}}' in p.text:
                    p.text = p.text.replace('{{PROJECT_CODE}}', project_code).replace('{{Project Code}}', project_code)
                    is_project_code = True
                
                if '{{REPORT_DATE}}' in p.text:
                    p.text = p.text.replace('{{REPORT_DATE}}', current_date_cover)
                    is_date = True
                
                if is_report_title and p.runs:
                    for run in p.runs:
                        run.font.name = 'Tahoma'
                        run.font.size = Pt(20)
                        run.font.bold = True
                elif is_project_code and p.runs:
                    for run in p.runs:
                        run.font.name = 'Tahoma'
                        run.font.size = Pt(15)
                        run.font.bold = True
                        p.paragraph_format.line_spacing = 2.5
                elif is_date and p.runs:
                    for run in p.runs:
                        run.font.name = 'Tahoma'
                        run.font.size = Pt(12)
                        run.font.bold = True
                        p.paragraph_format.line_spacing = 5.5

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            is_report_title = False
                            is_project_code = False
                            is_date = False

                            if '{{REPORT_TITLE}}' in p.text:
                                p.text = p.text.replace('{{REPORT_TITLE}}', report_title)
                                is_report_title = True
                            if '{{PROJECT_CODE}}' in p.text or '{{Project Code}}' in p.text:
                                p.text = p.text.replace('{{PROJECT_CODE}}', project_code).replace('{{Project Code}}', project_code)
                                is_project_code = True
                            
                            if '{{REPORT_DATE}}' in p.text:
                                p.text = p.text.replace('{{REPORT_DATE}}', current_date_cover)
                                is_date = True
                            
                            if is_report_title and p.runs:
                                for run in p.runs:
                                    run.font.name = 'Tahoma'
                                    run.font.size = Pt(20)
                                    run.font.bold = True
                            elif is_project_code and p.runs:
                                for run in p.runs:
                                    run.font.name = 'Tahoma'
                                    run.font.size = Pt(14)
                                    run.font.bold = True
                                
                            elif is_date and p.runs:
                                for run in p.runs:
                                    run.font.name = 'Tahoma'
                                    run.font.size = Pt(14)
                                    run.font.bold = True
                                    

            
            # Set margins to prevent overlapping with footer
            for section in doc.sections:
                section.bottom_margin = Inches(1.5)

            for tc_id, data in filtered_evidences.items():
                doc.add_page_break()
                
                # Info Table (Plain without blue backgrounds)
                table = doc.add_table(rows=7, cols=2)
                try:
                    table.style = 'Table Grid'
                except KeyError:
                    set_table_borders(table)
                table.allow_autofit = False
                table.columns[0].width = Inches(2.0)
                table.columns[1].width = Inches(4.5)
                
                rows = [
                    ("Test Case Id", tc_id),
                    ("Test Case Name", data["tc_name"]),
                    ("Date", current_date),
                    ("Status", data["status"]),
                    ("Test Steps", data.get("test_steps", "")),
                    ("Expected Result", data["expected_result"]),
                    ("Actual Result", data.get("actual_result", ""))
                ]
                
                for i, (label, val) in enumerate(rows):
                    cell_label = table.cell(i, 0)
                    cell_label.text = label
                    run_label = cell_label.paragraphs[0].runs[0]
                    run_label.bold = True
                    run_label.font.color.rgb = RGBColor(255, 255, 255)
                    set_cell_background(cell_label, '17375E')
                    
                    cell_val = table.cell(i, 1)
                    if (label == "Test Steps" or label == "Expected Result" or label == "Actual Result") and val:
                        # Properly render newlines in python-docx
                        cell_val.text = ""
                        for line in str(val).split('\n'):
                            if not cell_val.paragraphs:
                                p = cell_val.add_paragraph(line)
                            else:
                                if not cell_val.paragraphs[0].text:
                                    cell_val.paragraphs[0].text = line
                                else:
                                    cell_val.add_paragraph(line)
                    else:
                        cell_val.text = str(val)
                    if label == "Status":
                        run = cell_val.paragraphs[0].runs[0]
                        run.bold = True
                        if val.lower() == 'passed':
                            run.font.color.rgb = RGBColor(0, 128, 0)
                        else:
                            run.font.color.rgb = RGBColor(255, 0, 0)
                
                precondition = data.get("precondition", "1. Buka Postman\n2. Set request ke POST...")
                doc.add_paragraph(f"\nPrecondition:\n{precondition}")
                
                doc.add_paragraph()
                
                test_data_str = ""
                for api in data.get("api", []):
                    payload = api.get("request_payload", {})
                    if "nomor_transaksi" in payload and f"nomor_transaksi: {payload['nomor_transaksi']}" not in test_data_str:
                        test_data_str += f"nomor_transaksi: {payload['nomor_transaksi']}\n"
                    if "nomor_loan" in payload and f"nomor_loan: {payload['nomor_loan']}" not in test_data_str:
                        test_data_str += f"nomor_loan: {payload['nomor_loan']}\n"
                    if "ktp" in payload and f"ktp: {payload['ktp']}" not in test_data_str:
                        test_data_str += f"ktp: {payload['ktp']}\n"
                
                if not test_data_str and data.get("api"):
                    payload = data["api"][0].get("request_payload", {})
                    if payload:
                        test_data_str = json.dumps(payload, indent=2)

                if test_data_str:
                    p_td = doc.add_paragraph(f"Test Data:\n{test_data_str.strip()}")
                    p_td.paragraph_format.space_before = Pt(12)
                    p_td.paragraph_format.space_after = Pt(12)
                
                step_counter = 1
                for api in data["api"]:
                    endpoint_url = api.get('url', '')
                    endpoint_name = "API"
                    if "draft-akseptasi" in endpoint_url: endpoint_name = "Draft Akseptasi"
                    elif "inquiry" in endpoint_url: endpoint_name = "Inquiry Loan"
                    elif "otorisasi" in endpoint_url: endpoint_name = "Otorisasi"
                    elif "payment" in endpoint_url or "pembayaran" in endpoint_url: endpoint_name = "Payment/Pembayaran"
                    elif "batal" in endpoint_url or "cancel" in endpoint_url: endpoint_name = "Pembatalan"
                    elif "calculate" in endpoint_url or "kalkulator" in endpoint_url: endpoint_name = "Kalkulator Premi"
                    
                    p_step = doc.add_paragraph(f"\n[Test_Step_{step_counter}]: Hit API {endpoint_name} ({api['method']} {endpoint_url})", style='Heading 3')
                    p_step.paragraph_format.space_before = Pt(14)
                    p_step.paragraph_format.space_after = Pt(6)
                    
                    api_table = doc.add_table(rows=2, cols=2)
                    try:
                        api_table.style = 'Table Grid'
                    except KeyError:
                        set_table_borders(api_table)
                    api_table.allow_autofit = False
                    api_table.columns[0].width = Inches(3.25)
                    api_table.columns[1].width = Inches(3.25)
                    
                    req_cell = api_table.cell(0, 0)
                    req_cell.text = "Request"
                    req_run = req_cell.paragraphs[0].runs[0]
                    req_run.bold = True
                    req_run.font.color.rgb = RGBColor(255, 255, 255)
                    set_cell_background(req_cell, '17375E')
                    
                    res_cell = api_table.cell(0, 1)
                    res_cell.text = "Response"
                    res_run = res_cell.paragraphs[0].runs[0]
                    res_run.bold = True
                    res_run.font.color.rgb = RGBColor(255, 255, 255)
                    set_cell_background(res_cell, '17375E')
                    
                    req_text = json.dumps(api['request_payload'], indent=2)
                    res_text = json.dumps(api['response_json'], indent=2)
                    
                    api_table.cell(1, 0).text = req_text
                    api_table.cell(1, 1).text = res_text
                    
                    # Only apply Courier New to row 1 (the content)
                    for cell in api_table.rows[1].cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Courier New'
                                run.font.size = Pt(8)
                    step_counter += 1
                    
                for db in data["db"]:
                    p_db = doc.add_paragraph(f"\n[Test_Step_{step_counter}]: DB Validation", style='Heading 3')
                    p_db.paragraph_format.space_before = Pt(14)
                    p_db.paragraph_format.space_after = Pt(6)
                    doc.add_paragraph(f"Query: {db['query']}")
                    
                    if db['result'] and len(db['result']) > 0:
                        keys = list(db['result'][0].keys()) if isinstance(db['result'][0], dict) else range(len(db['result'][0]))
                        
                        max_cols = 5
                        chunks = [keys[i:i + max_cols] for i in range(0, len(keys), max_cols)]
                        
                        for chunk_idx, chunk_keys in enumerate(chunks):
                            if chunk_idx > 0:
                                doc.add_paragraph() # Spacing between tables
                            
                            db_table = doc.add_table(rows=1, cols=len(chunk_keys))
                            try:
                                db_table.style = 'Table Grid'
                            except KeyError:
                                set_table_borders(db_table)
                            db_table.allow_autofit = True
                            
                            for col_idx, key in enumerate(chunk_keys):
                                head_cell = db_table.cell(0, col_idx)
                                head_cell.text = str(key)
                                head_run = head_cell.paragraphs[0].runs[0]
                                head_run.bold = True
                                head_run.font.color.rgb = RGBColor(255, 255, 255)
                                set_cell_background(head_cell, '17375E')
                                
                            for row in db['result']:
                                row_cells = db_table.add_row().cells
                                for col_idx, key in enumerate(chunk_keys):
                                    val = row[key] if isinstance(row, dict) else row[col_idx]
                                    row_cells[col_idx].text = str(val)
                    else:
                        doc.add_paragraph("No records found.")
                        
                    step_counter += 1
                    
                if "epolis" in data:
                    epolis = data["epolis"]
                    p_epolis = doc.add_paragraph(f"\n[Test_Step_{step_counter}]: Pengecekan Scan QR & Lampiran E-Polis", style='Heading 3')
                    p_epolis.paragraph_format.space_before = Pt(14)
                    p_epolis.paragraph_format.space_after = Pt(6)
                
                    p_qr = doc.add_paragraph()
                    p_qr.add_run("Terjemahan QR Code:\n").bold = True
                    p_qr.add_run(epolis.get("qr_result", "N/A"))
                
                    for img_path in epolis.get("image_paths", []):
                        try:
                            doc.add_picture(img_path, width=Inches(5.0))
                            doc.add_paragraph() # Spacing between pages
                        except Exception as e:
                            doc.add_paragraph(f"[Gagal melampirkan gambar E-Polis: {e}]")
                        
                    step_counter += 1
                
                if "ui" in data:
                    p_ui = doc.add_paragraph(f"\n[Test_Step_{step_counter}]: Pengecekan UI (ACS/FMS)", style='Heading 3')
                    p_ui.paragraph_format.space_before = Pt(14)
                    p_ui.paragraph_format.space_after = Pt(6)
                    for ui_ev in data["ui"]:
                        sys_name = ui_ev.get("system_name", "System")
                        img_path = ui_ev.get("screenshot_path", "")
                    
                        p_sys = doc.add_paragraph()
                        p_sys.add_run(f"System: {sys_name}\n").bold = True
                    
                        try:
                            doc.add_picture(img_path, width=Inches(6.0))
                            doc.add_paragraph()
                        except Exception as e:
                            doc.add_paragraph(f"[Gagal melampirkan screenshot {sys_name}: {e}]")
                        
                    step_counter += 1
                    
                if "premi_validation" in data:
                    pv = data["premi_validation"]
                    p_pv = doc.add_paragraph(f"\n[Test_Step_{step_counter}]: Tabel Validasi Nilai Premi", style='Heading 3')
                    p_pv.paragraph_format.space_before = Pt(14)
                    p_pv.paragraph_format.space_after = Pt(6)
                    
                    pv_table = doc.add_table(rows=2, cols=4)
                    try:
                        pv_table.style = 'Table Grid'
                    except KeyError:
                        set_table_borders(pv_table)
                    pv_table.allow_autofit = True
                    
                    headers = ["Response Postman (API)", "Database", "UI ACS", "Status"]
                    for idx, h in enumerate(headers):
                        h_cell = pv_table.cell(0, idx)
                        h_cell.text = h
                        h_run = h_cell.paragraphs[0].runs[0]
                        h_run.bold = True
                        h_run.font.color.rgb = RGBColor(255, 255, 255)
                        set_cell_background(h_cell, '17375E')
                    
                    vals = [f"Rp {pv['api']:,.2f}", f"Rp {pv['db']:,.2f}", f"Rp {pv['acs']:,.2f}", pv['status']]
                    for idx, v in enumerate(vals):
                        v_cell = pv_table.cell(1, idx)
                        v_cell.text = v
                        if idx == 3:
                            v_run = v_cell.paragraphs[0].runs[0]
                            v_run.bold = True
                            if pv['status'] == "Passed":
                                v_run.font.color.rgb = RGBColor(0, 128, 0)
                            else:
                                v_run.font.color.rgb = RGBColor(255, 0, 0)
                                
                    step_counter += 1
            
            # Setup Page Number in Footer
            
            def setup_footer(footer_obj):
                # Clear any existing junk from footer
                for p in footer_obj.paragraphs:
                    p.text = ""
                    pPr = p._p.get_or_add_pPr()
                    pBdr = pPr.find(qn('w:pBdr'))
                    if pBdr is not None:
                        pPr.remove(pBdr)
                for t in footer_obj.tables:
                    t._element.getparent().remove(t._element)
                    
                # Create a 1-column small box on the right for page number
                footer_table = footer_obj.add_table(rows=1, cols=1, width=Inches(0.5))
                footer_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
                footer_table.autofit = False
                footer_table.columns[0].width = Inches(0.5)
                
                # Make the table borders invisible just in case
                tblPr = footer_table._element.tblPr
                
                cell = footer_table.cell(0, 0)
                set_cell_background(cell, '17375E')
                
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.font.size = Pt(11)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                add_page_number(run)

            for section in doc.sections:
                # Aktifkan pengaturan 'Different First Page' agar cover beda dengan isinya
                section.different_first_page_header_footer = True
                
                import copy
                # Copy header dari halaman utama ke cover agar logo muncul
                first_header = section.first_page_header
                for el in list(first_header._element):
                    first_header._element.remove(el)
                for el in section.header._element:
                    first_header._element.append(copy.deepcopy(el))
                
                # Copy relationships (seperti gambar logo) agar rId valid di header baru
                for rel_id, rel in section.header.part.rels.items():
                    if rel_id not in first_header.part.rels:
                        first_header.part.rels.add_relationship(
                            rel.reltype,
                            rel._target,
                            rel.rId,
                            rel.is_external
                        )
                    
                # Bersihkan footer di halaman pertama (Cover) agar tidak ada nomor halaman
                first_footer = section.first_page_footer
                for p in first_footer.paragraphs:
                    p.text = ""
                for t in first_footer.tables:
                    t._element.getparent().remove(t._element)
                    
                # Pasang kotak penomoran biru hanya untuk halaman utama & genap (mulai halaman 2)
                setup_footer(section.footer)
                setup_footer(section.even_page_footer)

            os.makedirs(os.path.dirname(output_file), exist_ok=True)
            doc.save(output_file)
            print(f"DOCX generated successfully: {output_file}")
        except Exception as e:
            print(f"Failed to generate DOCX: {e}")

    def generate_excel(self, evidences: dict, output_file: str):

        os.makedirs(os.path.dirname(output_file), exist_ok=True)
        try:
            wb = load_workbook('collections/test_script.xlsx')
            ws = wb.active
            
            # Fill down Stream and Nama Modul BEFORE deleting rows
            last_stream = None
            last_modul = None
            for row in range(4, ws.max_row + 1):
                stream_val = ws.cell(row=row, column=2).value
                modul_val = ws.cell(row=row, column=3).value
                if stream_val: last_stream = stream_val
                elif last_stream: ws.cell(row=row, column=2).value = last_stream
                if modul_val: last_modul = modul_val
                elif last_modul: ws.cell(row=row, column=3).value = last_modul

            # Hapus row yang tidak dieksekusi (dijalankan dari bawah ke atas agar index tidak bergeser)
            for row in range(ws.max_row, 3, -1):
                tc_id = ws.cell(row=row, column=4).value
                if not tc_id or not isinstance(tc_id, str) or not tc_id.startswith("TC-"):
                    continue
                if tc_id not in evidences:
                    ws.delete_rows(row, 1)

            # Update sisa row yang dieksekusi
            for row in range(4, ws.max_row + 1):
                tc_id = ws.cell(row=row, column=4).value
                if not tc_id or not isinstance(tc_id, str) or not tc_id.startswith("TC-"):
                    continue
                    
                if tc_id in evidences:
                    data = evidences[tc_id]
                    
                    if data.get("api") and len(data["api"]) > 0:
                        actual_json = data["api"][-1].get("response_json", {})
                        if actual_json:
                            ws.cell(row=row, column=12).value = "1. Response API:\n" + json.dumps(actual_json, indent=2)
                            ws.cell(row=row, column=12).alignment = Alignment(wrap_text=True, vertical="top")
                            
                    ws.cell(row=row, column=11).value = data.get("status", "Failed")
                    
                    if data.get("test_steps"):
                        ws.cell(row=row, column=8).value = data["test_steps"]
                        ws.cell(row=row, column=8).alignment = Alignment(wrap_text=True, vertical="top")
                    
                    
                    test_data_str = ""
                    for api in data.get("api", []):
                        payload = api.get("request_payload", {})
                        if "nomor_transaksi" in payload and f"nomor_transaksi: {payload['nomor_transaksi']}" not in test_data_str:
                            test_data_str += f"nomor_transaksi: {payload['nomor_transaksi']}\n"
                        if "nomor_loan" in payload and f"nomor_loan: {payload['nomor_loan']}" not in test_data_str:
                            test_data_str += f"nomor_loan: {payload['nomor_loan']}\n"
                        if "ktp" in payload and f"ktp: {payload['ktp']}" not in test_data_str:
                            test_data_str += f"ktp: {payload['ktp']}\n"
                    if test_data_str:
                        ws.cell(row=row, column=9).value = test_data_str.strip()
                        ws.cell(row=row, column=9).alignment = Alignment(wrap_text=True, vertical="top")
                        
                    ws.cell(row=row, column=13).value = "Fathur"
                    ws.cell(row=row, column=14).value = datetime.now().strftime("%d-%b-%Y")

            # Merge Title A1 to N2, center, wrap text, size 20, bold
            # Unmerge any existing ranges to prevent Excel corruption
            for range_str in list(ws.merged_cells.ranges):
                if "A1" in str(range_str):
                    try:
                        ws.unmerge_cells(str(range_str))
                    except Exception:
                        pass
            try:
                ws.merge_cells("A1:N2")
            except Exception:
                pass
            
            report_title = os.environ.get("REPORT_TITLE", "Laporan Hasil Pengujian API")
            project_code = os.environ.get("PROJECT_CODE", "PRJ-000")
            
            title_cell = ws.cell(row=1, column=1)
            title_cell.value = f"Test Script - {project_code} {report_title}"
            title_cell.font = Font(size=20, bold=True)
            title_cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

            # Format Header Row (Row 3): Blue background, White text, Bold
            header_fill = PatternFill(start_color="003366", end_color="003366", fill_type="solid")
            header_font = Font(color="FFFFFF", bold=True, size=16)
            for col in range(1, ws.max_column + 1):
                header_cell = ws.cell(row=3, column=col)
                header_cell.fill = header_fill
                header_cell.font = header_font
                header_cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                
            # Rename Notes (Column 12) to Actual Result
            ws.cell(row=3, column=12).value = "Actual Result"

            # Apply Full Border to the table (Row 3 to Max Row) and Wrap Text for data rows
            thin_border = Border(left=Side(style='thin'), 
                                 right=Side(style='thin'), 
                                 top=Side(style='thin'), 
                                 bottom=Side(style='thin'))
            
            for r in ws.iter_rows(min_row=3, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
                for cell in r:
                    cell.border = thin_border
                    if cell.row > 3:
                        cell.alignment = Alignment(wrap_text=True, vertical="top")

            # Auto-fit column widths
            for col_idx, col in enumerate(ws.columns, 1):
                max_length = 0
                col_letter = get_column_letter(col_idx)
                for cell in col:
                    if cell.row < 3:
                        continue
                    try:
                        if cell.value:
                            # Hitung panjang setiap baris dalam teks (karena ada newline)
                            lines = str(cell.value).split('\n')
                            longest_line = max(len(line) for line in lines)
                            if longest_line > max_length:
                                max_length = longest_line
                    except:
                        pass
                
                # Tambahkan sedikit padding
                adjusted_width = max_length + 2
                
                # Batasi lebar maksimal agar tidak terlalu lebar (misal max 50)
                if adjusted_width > 50:
                    adjusted_width = 50
                
                ws.column_dimensions[col_letter].width = adjusted_width

            # Save the final file to reports only (do not overwrite master test_script.xlsx)
            wb.save(output_file)
            
            print(f"Excel report generated successfully: {output_file}")
        except Exception as e:
            print(f"Failed to generate Excel report: {e}\n{traceback.format_exc()}")

    def get_next_defect_id(self):
        counter_file = "collections/defect_counter.json"
        count = 0
        if os.path.exists(counter_file):
            try:
                with open(counter_file, "r") as f:
                    data = json.load(f)
                    count = data.get("count", 0)
            except Exception:
                pass
        count += 1
        try:
            os.makedirs(os.path.dirname(counter_file), exist_ok=True)
            with open(counter_file, "w") as f:
                json.dump({"count": count}, f)
        except Exception as e:
            print(f"Failed to save defect counter: {e}")
        return f"DEF-{count}"

    def generate_defect_reports(self, evidences: dict, prefix: str, timestamp: str = None):
        failed_tcs = [tc_id for tc_id, data in evidences.items() if data.get("status", "").lower() == "failed"]
        if not failed_tcs:
            return

        try:
            import copy
            doc = Document('collections/TEMPLATE_Defect Report.docx')
            if not doc.tables:
                print("Template docx doesn't have tables")
                return

            template_tbl_element = copy.deepcopy(doc.tables[0]._element)
            is_first = True
            tracker_rows = []

            for tc_id in failed_tcs:
                data = evidences[tc_id]
                def_id = self.get_next_defect_id()
                
                if not is_first:
                    doc.add_page_break()
                    new_tbl = copy.deepcopy(template_tbl_element)
                    doc._body._element.append(new_tbl)
                    table = doc.tables[-1]
                else:
                    table = doc.tables[0]
                    is_first = False
                
                table.cell(0, 1).text = str(tc_id)
                table.cell(1, 1).text = def_id
                
                expected = data.get("expected_result", "Test execution failed.")
                table.cell(3, 1).text = f"Test execution failed.\n Expected: \n {expected}"
                
                cell_4_1 = table.cell(4, 1)
                cell_4_1.text = "" 
                
                current_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                p = cell_4_1.add_paragraph(f"Tanggal execution test: {current_time}\n")
                
                p.add_run("\nStep to reproduce:\n").bold = True
                
                step_counter = 1
                api_endpoints = []
                for api in data.get("api", []):
                    endpoint_url = api.get('url', '')
                    api_endpoints.append(endpoint_url)
                    endpoint_name = "API"
                    if "draft-akseptasi" in endpoint_url: endpoint_name = "Draft Akseptasi"
                    elif "inquiry" in endpoint_url: endpoint_name = "Inquiry Loan"
                    elif "otorisasi" in endpoint_url: endpoint_name = "Otorisasi"
                    elif "payment" in endpoint_url or "pembayaran" in endpoint_url: endpoint_name = "Payment/Pembayaran"
                    elif "batal" in endpoint_url or "cancel" in endpoint_url: endpoint_name = "Pembatalan"
                    elif "calculate" in endpoint_url or "kalkulator" in endpoint_url: endpoint_name = "Kalkulator Premi"
                    
                    p.add_run(f"{step_counter}. Hit {endpoint_name} ({api.get('method', 'POST')} {endpoint_url})\n")
                    step_counter += 1
                    
                for db in data.get("db", []):
                    p.add_run(f"{step_counter}. DB Validation\n")
                    step_counter += 1
                    
                if "ui" in data:
                    p.add_run(f"{step_counter}. Pengecekan UI ACS/FMS\n")
                    step_counter += 1
                    
                p.add_run("\nActual Result:\n").bold = True
                if data.get("api") and len(data["api"]) > 0:
                    last_api = data["api"][-1]
                    res_text = json.dumps(last_api.get('response_json', {}), indent=2)
                    p.add_run(f"API Response:\n{res_text}\n\n")
                
                if data.get("db") and len(data["db"]) > 0:
                    last_db = data["db"][-1]
                    db_results = last_db.get('result', [])
                    p.add_run("\nDB Result:\n").bold = True
                    
                    if db_results and isinstance(db_results, list) and isinstance(db_results[0], dict):
                        headers = list(db_results[0].keys())
                        db_table = cell_4_1.add_table(rows=1, cols=len(headers))
                        db_table.style = 'Table Grid'
                        
                        hdr_cells = db_table.rows[0].cells
                        for i, header in enumerate(headers):
                            hdr_cells[i].text = str(header)
                            for paragraph in hdr_cells[i].paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                                    
                        for row_data in db_results:
                            row_cells = db_table.add_row().cells
                            for i, header in enumerate(headers):
                                row_cells[i].text = str(row_data.get(header, ""))
                        
                        p = cell_4_1.add_paragraph("\n")
                    else:
                        db_res = json.dumps(db_results, indent=2, default=str)
                        p.add_run(f"{db_res}\n\n")

                if "epolis" in data:
                    epolis = data["epolis"]
                    p.add_run(f"Terjemahan QR Code: {epolis.get('qr_result', 'N/A')}\n\n")
                    for img_path in epolis.get("image_paths", []):
                        try:
                            p.add_run().add_picture(img_path, width=Inches(4.0))
                            p.add_run("\n")
                        except Exception as e:
                            p.add_run(f"[Gagal melampirkan gambar E-Polis: {e}]\n")
                            
                if "ui" in data:
                    for ui_ev in data["ui"]:
                        sys_name = ui_ev.get("system_name", "System")
                        img_path = ui_ev.get("screenshot_path", "")
                        p.add_run(f"System: {sys_name}\n")
                        try:
                            p.add_run().add_picture(img_path, width=Inches(4.0))
                            p.add_run("\n")
                        except Exception as e:
                            p.add_run(f"[Gagal melampirkan screenshot {sys_name}: {e}]\n")
                            
                # Collect for tracker
                endpoints_str = ", ".join(api_endpoints)
                actual_error_summary = "Failed during API execution"
                if "db" in data and len(data["db"]) > 0:
                    actual_error_summary = "Failed during DB Validation"
                if "ui" in data:
                    actual_error_summary = "Failed during UI Validation"
                    
                if data.get("api") and len(data["api"]) > 0:
                    last_api = data["api"][-1]
                    if last_api.get("response_status", 200) >= 400:
                        actual_error_summary = f"API Error {last_api.get('response_status')}"
                        
                tracker_rows.append([
                    def_id, 
                    current_time, 
                    str(tc_id), 
                    endpoints_str, 
                    str(expected), 
                    actual_error_summary
                ])
                            
            def setup_footer(footer_obj):
                for p_footer in footer_obj.paragraphs:
                    p_footer.text = ""
                    pPr = p_footer._p.get_or_add_pPr()
                    pBdr = pPr.find(qn('w:pBdr'))
                    if pBdr is not None:
                        pPr.remove(pBdr)
                for t in footer_obj.tables:
                    t._element.getparent().remove(t._element)
                    
                footer_table = footer_obj.add_table(rows=1, cols=1, width=Inches(0.5))
                footer_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
                footer_table.autofit = False
                footer_table.columns[0].width = Inches(0.5)
                
                cell = footer_table.cell(0, 0)
                set_cell_background(cell, '17375E')
                
                p_cell = cell.paragraphs[0]
                p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p_cell.add_run()
                run.font.size = Pt(11)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                add_page_number(run)

            for section in doc.sections:
                setup_footer(section.footer)
                        
            if timestamp:
                out_file = f"reports/defects/{prefix}Defect_Report_{timestamp}.docx"
            else:
                out_file = f"reports/defects/{prefix}Defect_Report.docx"
                
            os.makedirs(os.path.dirname(out_file), exist_ok=True)
            doc.save(out_file)
            print(f"Defect report generated successfully: {out_file}")
            
            # Update Master Tracker Excel
            if HAS_OPENPYXL and tracker_rows:
                from openpyxl import Workbook, load_workbook
                tracker_file = "reports/defects/Master_Defect_Tracker.xlsx"
                os.makedirs(os.path.dirname(tracker_file), exist_ok=True)
                
                if os.path.exists(tracker_file):
                    wb = load_workbook(tracker_file)
                    ws = wb.active
                else:
                    wb = Workbook()
                    ws = wb.active
                    ws.title = "Defect Tracker"
                    headers = ["Defect ID", "Timestamp", "Test Case ID", "API Endpoint", "Expected Result", "Actual Error Summary"]
                    ws.append(headers)
                    from openpyxl.styles import Font
                    for cell in ws[1]:
                        cell.font = Font(bold=True)
                        
                for row in tracker_rows:
                    ws.append(row)
                    
                wb.save(tracker_file)
                print(f"Master Defect Tracker updated: {tracker_file}")
            
        except Exception as e:
            print(f"Failed to generate merged defect report: {e}\n{traceback.format_exc()}")

report_generator = ReportGenerator()
